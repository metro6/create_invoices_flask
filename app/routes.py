import io
import os
import uuid
from datetime import datetime
from flask import render_template, flash, redirect, url_for, request, Response, send_file, send_from_directory
from flask_login import current_user, login_user, logout_user, login_required
from werkzeug.urls import url_parse
from werkzeug.utils import secure_filename


from docx import *
from docx.shared import Inches

from app import app, db

from .forms import LoginForm, RegistrationForm, CreateFormForm, EditFormForm
from .models import Invoice, User, Image



@app.route('/')
@app.route('/index')
@login_required
def index():
    invoices = Invoice.query.all()
    return render_template('index.html',
                           title='Список форм',
                           invoices=invoices)



@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user is None or not user.check_password(form.password.data):
            flash('Учетные данные введены неверно')
            return redirect(url_for('login'))
        login_user(user, remember=form.remember_me.data)
        next_page = request.args.get('next')
        if not next_page or url_parse(next_page).netloc != '':
            next_page = url_for('index')
        return redirect(next_page)
    return render_template('login.html',
                           title='Авторизация',
                           form=form)


@app.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('index'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = RegistrationForm()
    if form.validate_on_submit():
        user = User(username=form.username.data)
        user.set_password(form.password.data)
        db.session.add(user)
        db.session.commit()
        flash('Вы успешно зарегестрировались')
        return redirect(url_for('login'))
    return render_template('register.html',
                           title='Регистрация',
                           form=form)


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1] in app.config['ALLOWED_EXTENSIONS']


@app.route('/invoice/<invoice_id>', methods=['GET', 'POST'])
@login_required
def invoice(invoice_id):
    invoice = Invoice.query.filter_by(id=invoice_id).first_or_404()
    image = Image.query.filter_by(invoice_id=invoice.id).first()
    title = invoice.name
    form = EditFormForm()
    if form.validate_on_submit():
        invoice.name = form.name.data
        invoice.text = form.text.data
        invoice.full_text = form.full_text.data
        if form.departure_date.data:
            invoice.departure_date = datetime.strptime(form.departure_date.data, '%Y-%m-%d %H:%M:%S')
        else:
            invoice.departure_date = None
        if form.receive_date.data:
            invoice.receive_date = datetime.strptime(form.receive_date.data, '%Y-%m-%d %H:%M:%S')
        else:
            invoice.receive_date = None
        if form.delete_picture:
            img = Image.query.filter_by(invoice_id=invoice.id).first()
            if img:
                db.session.delete(img)
        if form.picture:
            file = request.files['picture']
            if file:
                filename = str(uuid.uuid4()) + '_' + secure_filename(file.filename)
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                img = Image.query.filter_by(invoice_id=invoice.id).first()
                if not img:
                    img = Image(path="/media/" + filename, invoice_id=invoice.id)
                else:
                    img.path = '/media/' + filename
                db.session.add(img)
        db.session.add(invoice)
        db.session.commit()
        return redirect('/invoice/' + invoice_id)

    elif request.method == 'GET':
        form.name.data = invoice.name
        form.text.data = invoice.text
        form.full_text.data = invoice.full_text
        form.departure_date.data = str(invoice.departure_date)
        form.receive_date.data = str(invoice.receive_date)
    return render_template('invoice.html',
                           title=title,
                           invoice=invoice,
                           form=form,
                           image=image)


@app.route('/create_form', methods=['GET', 'POST'])
@login_required
def create_form():
    form = CreateFormForm()
    title = 'Создание новой формы'
    if form.validate_on_submit():
        invoice = Invoice(name=form.name.data)
        db.session.add(invoice)
        db.session.commit()
        flash('Новая форма успешно создана')
        return redirect(url_for('index'))
    return render_template('create_form.html',
                           title=title,
                           form=form)


@app.route('/media/<filename>')
def media(filename):
    return send_from_directory('media', filename)


@app.route('/get_invoice_word')
def get_invoice_word():
    invoice = Invoice.query.filter_by(id=request.args['id']).first_or_404()
    image = Image.query.filter_by(invoice_id=invoice.id).first()
    print('@@@@@@@@@@@')

    full_text = request.args['full'] == 'true'
    short_text = request.args['short'] == 'true'

    file_stream = io.BytesIO()

    document = Document()
    document.add_heading(invoice.name, 0)
    document.add_paragraph('Дата создания документа ' + datetime.now().strftime('%d-%m-%Y %H:%M:%S'))

    if image:
        document.add_picture(os.path.join(app.root_path) + image.path, width=Inches(1.25))

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название поля'
    hdr_cells[1].text = 'Значение поля'

    records = []

    if full_text:
        if invoice.full_text:
            records.append(['Полное описание', invoice.full_text])

    if short_text:
        if invoice.text:
            records.append(['Краткое описание', invoice.text])

    if invoice.departure_date:
        records.append(['Дата отправления', invoice.departure_date.strftime('%d-%m-%Y %H:%M:%S')])

    if invoice.receive_date:
        records.append(['Дата получения', invoice.receive_date.strftime('%d-%m-%Y %H:%M:%S')])

    for key, value in records:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value

    document.add_page_break()
    document.save(file_stream)

    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, attachment_filename='report_.docx')